#!/usr/bin/env python3
"""
Smart SOV parser for Construction AI Billing.
Accepts PDF (.pdf) and Word (.docx/.doc) contractor estimates.
Extracts: item_id, description, scheduled_value
Returns JSON to stdout with separate summary section.
"""
import sys, json, re, os

# ── Patterns ───────────────────────────────────────────────────────────────────
# Pattern to match financial summary lines (subtotal, total, overhead, tax, etc.)
SUMMARY_RE = re.compile(
    r'^(subtotal|company\s+overhead|overhead|tax\s+state\s+tax[^$]*|tax|state\s+tax|balance\s+due|amount\s+paid|total)'
    r'(?:\s.*)?$',
    re.IGNORECASE
)

# Pattern to skip non-financial metadata (headers, footers, terms, signatures, etc.)
SKIP_RE = re.compile(
    r'^(\*|•|·|–|—|-{2,})'               # bullet / sub-item lines
    r'|^(terms\s+and\s+conditions|signature|page \d+\s*/|files\+|note[:\s]|excludes'
    r'|it is an honor|we thank|sincerely|dear |http|www\.'
    r'|material and\s+labor|item\s+labor|labor\s+overhead)',
    re.IGNORECASE
)

def clean_desc(s):
    s = re.sub(r'^[\*\•\-–—·]+\s*', '', s.strip())   # remove leading bullets
    s = re.sub(r'\s+', ' ', s).strip()
    return s

def clean_text(text):
    """Remove null characters and other problematic control characters."""
    return text.replace('\x00', '').strip()

def extract_amounts(text):
    """Return list of floats from all $X,XXX.XX patterns in text."""
    text = clean_text(text)
    return [float(m.replace('$','').replace(',',''))
            for m in re.findall(r'\$[\d,]+(?:\.\d{1,2})?', text)]

def extract_amounts_loose(text):
    """Return list of floats from $X,XXX.XX OR bare X,XXX.XX patterns (for summary lines)."""
    text = clean_text(text)
    # First try with $ prefix
    dollar = [float(m.replace('$','').replace(',',''))
              for m in re.findall(r'\$[\d,]+(?:\.\d{1,2})?', text)]
    if dollar:
        return dollar
    # Fallback: bare comma-formatted numbers (e.g. "201,186.41")
    bare = [float(m.replace(',',''))
            for m in re.findall(r'\b\d{1,3}(?:,\d{3})+(?:\.\d{1,2})?\b', text)]
    return bare

def assign_id(counter, existing_id=''):
    """Return existing_id if it looks like a real cost code, else auto-assign."""
    if re.match(r'^\d{3,6}$', str(existing_id).strip()):
        return str(existing_id).strip()
    val = str(counter[0])
    counter[0] += 1000
    return val

def extract_summary_label(line):
    """Extract and normalize summary label from a line.
    Returns (key, full_label) or (None, None)."""
    line_stripped = line.strip()
    match = SUMMARY_RE.match(line_stripped)
    if not match:
        return None, None
    label = match.group(1).strip().lower()
    # Normalize labels (order matters — check specific before generic)
    if 'subtotal' in label:
        return 'subtotal', line_stripped
    elif 'overhead' in label:
        return 'overhead', line_stripped
    elif 'tax' in label:
        return 'tax', line_stripped
    elif 'balance' in label and 'due' in label:
        return 'balance_due', line_stripped
    elif 'amount' in label and 'paid' in label:
        return 'amount_paid', line_stripped
    elif 'total' in label:
        return 'total', line_stripped
    return None, None

# ── PDF parser ─────────────────────────────────────────────────────────────────
def parse_pdf(filepath):
    import pdfplumber
    rows = []
    summary = {}
    counter = [1000]
    seen = set()

    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ''
            text = clean_text(text)
            for line in text.split('\n'):
                line = clean_text(line)
                if len(line) < 5:
                    continue

                # Check if this is a summary line
                summary_key, full_label = extract_summary_label(line)
                if summary_key:
                    amounts = extract_amounts_loose(line)
                    if amounts:
                        summary[summary_key] = round(amounts[-1], 2)
                        if summary_key == 'tax':
                            # Capture the full tax description e.g. "Tax State tax(7.25%)"
                            tax_part = re.sub(r'\$?[\d,]+(?:\.\d{1,2})?', '', line).strip()
                            tax_part = re.sub(r'\s+', ' ', tax_part).strip()
                            if tax_part:
                                summary['tax_label'] = tax_part
                    continue

                if SKIP_RE.search(line):
                    continue

                amounts = extract_amounts(line)
                if not amounts:
                    continue

                # Last dollar amount = Total column (Material + Labor + Overhead)
                total = amounts[-1]
                if total <= 0:
                    continue

                # Description = everything before the first dollar sign
                desc = re.sub(r'\s*\$[\d,]+(?:\.\d{1,2})?.*$', '', line).strip()
                desc = clean_desc(desc)

                # Skip very short or numeric-only descriptions
                if len(desc) < 4 or re.match(r'^[\d\s\.\-]+$', desc):
                    continue
                if SKIP_RE.search(desc):
                    continue

                key = desc.lower()
                if key in seen:
                    continue
                seen.add(key)

                rows.append({
                    'item_id': assign_id(counter),
                    'description': desc,
                    'scheduled_value': round(total, 2)
                })

    return rows, summary

# ── DOCX parser ────────────────────────────────────────────────────────────────
def parse_docx(filepath):
    import docx as python_docx
    doc = python_docx.Document(filepath)
    rows = []
    summary = {}
    counter = [1000]
    seen = set()

    def add_row(item_id, desc, amt):
        desc = clean_desc(desc)
        if len(desc) < 4 or SKIP_RE.search(desc):
            return
        key = desc.lower()
        if key in seen or amt <= 0:
            return
        seen.add(key)
        rows.append({
            'item_id': assign_id(counter, item_id),
            'description': desc,
            'scheduled_value': round(amt, 2)
        })

    # Tables first (most structured)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if not any(cells):
                continue

            # Check if any cell is a summary line
            found_summary = False
            for ci, cell_text in enumerate(cells):
                summary_key, full_label = extract_summary_label(cell_text)
                if summary_key:
                    # Look for amount in this cell or the rightmost cell with a number
                    amts = extract_amounts_loose(cell_text)
                    if not amts:
                        # Try other cells in this row for the amount
                        for cj in range(len(cells) - 1, -1, -1):
                            amts = extract_amounts_loose(cells[cj])
                            if amts:
                                break
                    if amts:
                        summary[summary_key] = round(amts[-1], 2)
                        if summary_key == 'tax' and 'tax' in cell_text.lower():
                            tax_part = re.sub(r'\$?[\d,]+(?:\.\d{1,2})?', '', cell_text).strip()
                            if tax_part:
                                summary['tax_label'] = tax_part
                    found_summary = True
                    break

            if found_summary:
                continue

            # Find rightmost numeric cell as amount; leftmost text as desc
            amt_idx = -1
            for ci in range(len(cells) - 1, -1, -1):
                amts = extract_amounts(cells[ci])
                if amts and amts[-1] > 0:
                    amt_idx = ci
                    break
            if amt_idx < 0:
                continue
            # Description = longest non-numeric cell to the left of amount
            desc_candidates = [c for c in cells[:amt_idx] if len(c) > 3 and not extract_amounts(c)]
            if not desc_candidates:
                continue
            desc = max(desc_candidates, key=len)
            item_id = cells[0] if re.match(r'^\d{3,6}$', cells[0]) else ''
            add_row(item_id, desc, extract_amounts(cells[amt_idx])[-1])

    # Fallback: paragraphs
    if not rows:
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text or SKIP_RE.search(text):
                continue

            # Check if this is a summary line
            summary_key, full_label = extract_summary_label(text)
            if summary_key:
                amts = extract_amounts_loose(text)
                if amts:
                    summary[summary_key] = round(amts[-1], 2)
                    if summary_key == 'tax' and 'tax' in text.lower():
                        tax_part = re.sub(r'\$?[\d,]+(?:\.\d{1,2})?', '', text).strip()
                        if tax_part:
                            summary['tax_label'] = tax_part
                continue

            amts = extract_amounts(text)
            if not amts:
                continue
            total = amts[-1]
            desc = re.sub(r'\s*\$[\d,]+(?:\.\d{1,2})?.*$', '', text).strip()
            add_row('', desc, total)

    return rows, summary

# ── Main ───────────────────────────────────────────────────────────────────────
def main():
    if len(sys.argv) < 2:
        print(json.dumps({'error': 'No file path provided'}))
        sys.exit(1)

    filepath = sys.argv[1]
    ext = os.path.splitext(filepath)[1].lower()

    try:
        if ext == '.pdf':
            rows, summary = parse_pdf(filepath)
        elif ext in ('.docx', '.doc'):
            rows, summary = parse_docx(filepath)
        else:
            print(json.dumps({'error': f'Unsupported format: {ext}. Use PDF or DOCX.'}))
            sys.exit(1)

        if not rows:
            print(json.dumps({'error': 'No line items with dollar amounts found in this file.'}))
            sys.exit(1)

        # Compute total from line items
        computed_total = sum(row['scheduled_value'] for row in rows)

        print(json.dumps({
            'rows':            rows,
            'all_rows':        rows,
            'row_count':       len(rows),
            'total_rows':      len(rows),
            'computed_total':  round(computed_total, 2),
            'summary':         summary,
            'sheet_used':      ext.lstrip('.').upper(),
            'filename':        os.path.basename(filepath)
        }))

    except Exception as e:
        print(json.dumps({'error': str(e)}))
        sys.exit(1)

if __name__ == '__main__':
    main()
